
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor, Inches, Pt
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from matplotlib.backends.backend_pdf import PdfPages
from datetime import datetime

import matplotlib.pyplot as plt
import pandas as pd
import tabula
import string
import docx
import os

pd.options.display.max_columns = None  # display options for table
pd.options.display.width = None  # allows 'print table' to fill output screen
pd.options.mode.chained_assignment = None  # disables error caused by chained dataframe iteration


# input: elevation photos and xrf pos photos
# output: photo log as docx and pdf
def create_photo_log(beholden):
    doc = docx.Document()  # create instance of a word document
    sections = doc.sections  # change the page margins
    for section in sections:  # set margins equal to 0 on all sides of doc container
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # create elevation and xrf photo dic
    f_name = str(beholden[0]) + ' - ' + str(beholden[5] + ' - ' + str(beholden[6]))

    elev_lab_lis = []
    for x in range(4):
        elev_lab_lis.append('Elevation ' + string.ascii_uppercase[x])
    arr_elev_pat = []
    for x in range(4):
        arr_elev_pat.append(str('job_Folders/' + f_name + '/' + str(beholden[2]) + '_LBP/elevations/' + string.ascii_lowercase[x] + '.jpg'))

    xrf_lab_lis = []
    xrf_pos_pat_lis = os.listdir('job_Folders/' + f_name + '/' + str(beholden[2]) + '_LBP/xrf_Photos')
    for x in range(len(xrf_pos_pat_lis)):
        xrf_lab_lis.append('Reading ' + str(xrf_pos_pat_lis[x].split('_')[1]))
    arr_xrf_pat = []
    for x in range(len(xrf_lab_lis)):
        numhold = xrf_lab_lis[x].split(' ')[1]
        arr_xrf_pat.append('job_Folders/' + f_name + '/' + str(beholden[2]) + '_LBP/xrf_Photos/' + str([num for num in xrf_pos_pat_lis if str(numhold) in str(num)][0]))

    lab_full = elev_lab_lis + xrf_lab_lis
    pat_full = arr_elev_pat + arr_xrf_pat
    page_len = round(len(lab_full) / 6)
    label_groups = [lab_full[i:i + 6] for i in range(0, len(lab_full), 6)]
    pat_groups = [pat_full[i:i + 6] for i in range(0, len(pat_full), 6)]

    table_arr = []
    header_arr = []
    brk_para_arr = []
    brk_run_arr = []
    # ------------------------------------------------------------------------------------------------------------------
    for x in range(page_len):
        # add header
        header_arr.append(doc.add_table(1, 2))
        header_widths = [7, 3]
        header_arr[x].alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(1, 2):  # set table cell widths
            for cell in header_arr[x].columns[i].cells:
                cell.width = Inches(header_widths[i])
                cell.height = Inches(1.5)

        # left cell
        tab_head_left_cell = header_arr[x].cell(0, 0)
        par_head_left = tab_head_left_cell.paragraphs[0]
        run_head_left = par_head_left.add_run('Photo Log - ' + beholden[0])
        font_head = run_head_left.font
        font_head.size = Pt(14)
        run_head_left.font.color.rgb = RGBColor(0, 91, 184)
        par_head_left.alignment = 0

        # right cell
        tab_head_right_cell = header_arr[x].cell(0, 1)
        par_head_right = tab_head_right_cell.paragraphs[0]
        run_head_right = par_head_right.add_run()
        run_head_right.add_picture('lead_Pit/reporting/photo_Log/ei.jpg', width=Inches(2))
        par_head_right.alignment = 2
        par_head_right.add_run()

        # t_head formatting
        header_arr[x].cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # ------------------------------------------------------------------------------------------------------------------

        lg = label_groups[x]
        pg = pat_groups[x]
        table_arr.append(doc.add_table(len(lg), 2))
        table_arr[x].alignment = WD_TABLE_ALIGNMENT.CENTER
        table_arr[x].style = 'Table Grid'
        cell_arr = []
        par_arr = []
        run_arr = []
        for i in range(len(pg))[::2]:
            for j in range(2):
                ipj = i + j
                cell_arr.append(table_arr[x].cell(i, j))
                cell_arr[ipj].width = Inches(3.5)
                par_arr.append(cell_arr[ipj].paragraphs[0])
                # par_arr[ipj].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
                run_arr.append(par_arr[ipj].add_run())
                run_arr[ipj].add_picture(pg[ipj], height=Inches(2.75))
                last_p = doc.tables[-1].rows[-1].cells[-1].paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_label_arr = []
        par_label_arr = []
        run_label_arr = []
        font_label_arr = []
        for i in range(len(pg))[1::2]:
            for j in range(2):
                comb = i + j -1
                cell_label_arr.append(table_arr[x].cell(i, j))
                par_label_arr.append(cell_label_arr[comb].paragraphs[0])
                par_label_arr[comb].alignment = 1
                run_label_arr.append(par_label_arr[comb].add_run(lg[comb]))
                font_label_arr.append(run_label_arr[comb].font)
                font_label_arr[comb].size = Pt(14)

        skipper = 0
        for row in table_arr[x].rows:
            if skipper == 1:
                row.height = Pt(20)
                skipper = 0
            else:
                skipper = 1
        if page_len > 1 and x != (page_len-1):
            brk_para_arr.append(doc.add_paragraph())
            brk_run_arr.append(brk_para_arr[x].add_run(''))

    # ------------------------------------------------------------------------------------------------------------------
    # SAVE DOCUMENT AS

    fl_pat = r'C:\Users\Elliott\pythonplay\lead_Pit\LRA\finished_Docs'
    doc.save(str(fl_pat + '\\' + beholden[0] + '\\' + beholden[0] + '_photo_Log.docx'))
    # ------------------------------------------------------------------------------------------------------------------


# input: dfs is list of xrf pos tables
# input: findings is the blank table established in doc
# input: index is the table index in dfs
# output: populated and formatted table in doc
def pop_table(dfs, findings, index):
    dfs[index].loc[-1] = dfs[index].columns  # adds column row as index -1
    dfs[index].index = dfs[index].index + 1  # adds 1 to all indices, setting column header to index 0
    dfs[index].sort_index(inplace=True)  # sorts header to top of dataframe

    nonelis = []
    for x in range(dfs[index].shape[1]):
        nonelis.append(None)
    if index < 3:
        dfs[index].loc[-1] = nonelis  # placeholder gets replaced later
        t1_widths = [0.75, 0.5, 1.5, 1, 1.25, 1]
    if index == 5:
        dfs[5].loc[-1] = nonelis  # placeholder gets replaced later
        t1_widths = [1.25, 0.75, 1.75, 2]
    if index == 3 or index == 4:
        dfs[index].loc[-1] = nonelis  # placeholder gets replaced later
        t1_widths = [.5, 1.25, 1.5, 1.75, 1]

    dfs[index].index = dfs[index].index + 1
    dfs[index].sort_index(inplace=True)
    findings.alignment = WD_TABLE_ALIGNMENT.CENTER  # align table center

    for i in range(1, dfs[index].shape[1]):  # set column widths using t1_widths
        for cell in findings.columns[i].cells:
            cell.width = Inches(t1_widths[i])
            cell.height = Pt(5)
    table_names = ['Table 1: Lead-Based Paint¹',
                   'Table 2: Deteriorated Lead-Based Paint¹',
                   'Table 3: Lead Containing Materials²',
                   'Table 4: Dust Wipe Sample Analysis',
                   'Table 5: Soil Sample Analysis',
                   'Table 6: Lead Hazard Control Options¹']
    for i in range(dfs[index].shape[0]):  # range is height of df[index]
        if i == 0:  # fill cell and format header
            my_cell = findings.cell(0, 0)
            sml_p = my_cell.paragraphs[0]
            runp = sml_p.add_run('\n')
            font = runp.font
            font.size = Pt(4)
            my_paragraph = my_cell.paragraphs[0]
            my_paragraph.paragraph_format.line_spacing = 1.5
            run = my_paragraph.add_run(table_names[index])
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="064E91"/>'.format(nsdecls('w')))
            my_cell._tc.get_or_add_tcPr().append(shading_elm)
            my_paragraph.alignment = 1
        elif i == 1:  # fill cells and format column headers
            for j in range(dfs[index].shape[1]):
                my_cell = findings.cell(1, j)
                sml_p = my_cell.paragraphs[0]
                runp = sml_p.add_run('\n')
                font = runp.font
                font.size = Pt(4)

                my_paragraph = my_cell.paragraphs[0]
                my_paragraph.paragraph_format.line_spacing = 1.5
                run = my_paragraph.add_run(str(dfs[index].values[1, j]))
                run.font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = parse_xml(r'<w:shd {} w:fill="064E91"/>'.format(nsdecls('w')))
                my_cell._tc.get_or_add_tcPr().append(shading_elm)
                my_paragraph.alignment = 1
        else:  # fill cells and format body of table
            for j in range(dfs[index].shape[1]):
                my_cell = findings.cell(i, j)
                sml_p = my_cell.paragraphs[0]
                runp = sml_p.add_run('\n')
                font = runp.font
                font.size = Pt(1)

                my_paragraph = my_cell.paragraphs[0]
                my_paragraph.paragraph_format.line_spacing = 1.25
                run = my_paragraph.add_run(str(dfs[index].values[i, j]))
                my_paragraph.alignment = 1  # fill tables

    if index < 3:
        findings.cell(0, 0).merge(findings.cell(0, 5))
    if index == 5:
        findings.cell(0, 0).merge(findings.cell(0, 3))
    else:
        findings.cell(0, 0).merge(findings.cell(0, 4))
    findings.style = 'Table Grid'


# input: xrf clean
# output: dflis is a dictionary of xrf pos dataframes
def xrf_tables(xrf, pdf_path):

    pb_res = pdf_scrape(pdf_path)

    # ------------------------------------------------------------------------------------------------------------------
    # manipulating dataframes

    xrf_pos = xrf  # create dataframe named xrf_pos containing all non-calibration positive lead readings

    for index, row in xrf_pos.iterrows():  # remove all negative and calibration rows
        if xrf_pos.loc[index][8] == 'Negative' or xrf_pos.loc[index][2] == 'Calibration':
            xrf_pos = xrf_pos.drop(index)

    xrf_pos = pd.concat([xrf_pos['Room'],
                         xrf_pos['Side'],
                         xrf_pos['Component'],
                         xrf_pos['Substrate'],
                         xrf_pos['Condition'],
                         xrf_pos['Color']], axis=1)  # create subdf, xrf_pos, from df xrf_pos

    # ------------------------------------------------------------------------------------------------------------------
    # create df Table 1: Lead-Based Paint
    xrf_pos1 = xrf_pos.reset_index(drop=True)  # reset indices
    xrf_pos1 = xrf_pos1.rename(columns={'Component': 'Component²'})  # rename column name
    new_row = pd.DataFrame(list(xrf_pos1.columns)).T  # transform vertical initial dataframe to horizontal
    new_row.columns = list(xrf_pos1.columns)  # set array new_row equal to column names
    xrf_pos1 = xrf_pos1.reset_index(drop=True)  # reset indices

    # ------------------------------------------------------------------------------------------------------------------
    # create df "Table 2: Deteriorated Lead-Based Paint¹"
    xrf_pos2 = xrf_pos1
    for index, row in xrf_pos2.iterrows():  # drop all rows with intact paint
        if xrf_pos2.loc[index]['Condition'] == 'Intact':
            xrf_pos2 = xrf_pos2.drop(index)

    # ------------------------------------------------------------------------------------------------------------------
    # create df "Table 3, Lead Containing Materials²
    xrf_pos3 = xrf_pos1
    for index, row in xrf_pos3.iterrows():  # drop all rows not containing ceramic
        if xrf_pos3.loc[index]['Substrate'] != 'Ceramic':
            xrf_pos3 = xrf_pos3.drop(index)

    # ------------------------------------------------------------------------------------------------------------------
    # create df "Table 4, Dust Wipe Sample Analysis"

    pb_wipes = pb_res[0][0]
    pb_wipes = pb_wipes.drop(0)
    pb_wipes.rename(columns={pb_wipes.columns[1]: 'sample id'})
    pb_wipes1 = pb_wipes.iloc[::2]  # select every other row
    pb_wipes1.columns = ['Sample #', 'Location', 'Surface Type', 'Concentration (ug/ft²)', 'Lead Hazard¹', 'one', 'two']
    pb_wipes1 = pb_wipes1.reset_index(drop=True)
    pb_wipes2 = pb_wipes.iloc[1::2]  # select opposite other rows
    pb_wipes2.columns = ['one', 'two', 'three', 'four', 'five', 'six', 'seven']
    pb_wipes2 = pb_wipes2.reset_index(drop=True)
    pb_wipes2['six'] = pb_wipes2['six'].replace(['nan'], ['<5.00 μg/ft2'])

    pb_wipes3 = pb_wipes1
    for index, row in pb_wipes3.iterrows():
        pb_wipes1.at[index, 'Sample #'] = 'DW' + str(index + 1)
        pb_wipes3.at[index, 'Location'] = str(pb_wipes3.iloc[index]['Surface Type']).split(' ')[0]
        pb_wipes3.at[index, 'Surface Type'] = str(pb_wipes3.iloc[index]['Surface Type']).split(' ')[-1]
        pb_wipes3.at[index, 'Concentration (ug/ft²)'] = pb_wipes2.at[index, 'six']
        if list(str(pb_wipes2.at[index, 'six']))[0] == '<' or str(pb_wipes2.at[index, 'six']) == 'nan':
            pb_wipes1.at[index, 'Lead Hazard¹'] = 'No'
        else:
            pb_wipes1.at[index, 'Lead Hazard¹'] = 'Yes'

    pb_wipes3['Location'] = pb_wipes3['Location'].replace(
        ['Bath', 'Bed', 'QC', 'Field', 'LR', 'Living', 'BR', 'Fam'],
        ['Bathroom', 'Bedroom', 'Q/C', 'Q/C', 'Living Room', 'Living Room', 'Bedroom', 'Family Room'])

    pb_wipes3['Surface Type'] = pb_wipes3['Surface Type'].replace(
        ['QC', 'Q/C', 'Blank', 'Sill'],
        ['Blank Wipe', 'Blank Wipe', 'Blank Wipe', 'Windowsill'])

    pb_wipes3 = pb_wipes3.drop(['one', 'two'], axis=1)
    pb_wipes3['Concentration (ug/ft²)'] = pb_wipes3['Concentration (ug/ft²)'].map(lambda y: str(y).split(' ')[0])
    pb_wipes3['Concentration (ug/ft²)'] = pb_wipes3['Concentration (ug/ft²)'].replace(['nan'], ['<5.00 μg/wipe'])

    # ------------------------------------------------------------------------------------------------------------------
    # create df "Table 5: Soil Sample Analysis"

    pb_drip = pb_res[1][0]
    pb_drip.columns = ['Sample #', 'Location', 'Bare/Covered', 'Concentration (mg/kg)', 'Lead Hazard¹', 'one', 'two',
                       'three']
    pb_drip = pb_drip.drop([0], axis=0)

    if list(pb_drip.at[2, 'two'])[0] == '<':
        pb_haz = 'No'
    else:
        pb_haz = 'Yes'

    drip_row = ['S-1', str(pb_drip.at[1, 'Bare/Covered']).split(' ')[-1], 'Bare', pb_drip.at[2, 'two'], pb_haz, 'hold',
                'hold', 'hold']
    pb_drip.loc[1] = drip_row
    pb_drip = pb_drip.drop(2)
    pb_drip.at[1, 'Sample #'] = 'S-1'
    pb_drip.at[1, 'Concentration (mg/kg)'] = str(pb_drip.at[1, 'Concentration (mg/kg)']).split(' ')[0]
    pb_drip = pb_drip.drop(['one', 'two', 'three'], axis=1)

    # ------------------------------------------------------------------------------------------------------------------
    # create df "Table 6: Lead Hazard Control Options¹"

    tab6 = pb_drip.drop([1], axis=0)  # copy blank table 5 to tab6, drop last row
    tab6.drop('Lead Hazard¹', axis=1, inplace=True)
    tab6.columns = ['Hazard Type', 'Location', 'Description', 'Control²⁻⁵']  # set columns of tab6

    ldhlis = []  # blank list to hold lead hazard locations
    ldsurlis = []  # blank list to hold lead hazard surface types
    for index, row in pb_wipes3.iterrows():
        if pb_wipes3.at[index, 'Lead Hazard¹'] == 'Yes':
            ldhlis.append(pb_wipes3.at[index, 'Location'])
            ldsurlis.append(pb_wipes3.at[index, 'Surface Type'])

    if pb_wipes3['Lead Hazard¹'].str.contains('Yes').any():
        shold = [['Lead Dust Hazard',
                  str(unq_lis(ldhlis)),
                  str(unq_lis(ldsurlis)),
                  'Cleaning- Clean surfaces using HEPA filtered vacuum and wet cleaning agents to '
                  'remove leaded dust']]
        s4 = pd.DataFrame(shold, columns=tab6.columns)
        tab6 = pd.concat([tab6, s4], axis=0)

    if pb_drip['Lead Hazard¹'].str.contains('Yes').any():
        shold = [['Lead Soil Hazard',
                  'Exterior',
                  'Drip Line',
                  'Abate and remove lead-contaminated soil or overlay a durable covering such as '
                  'asphalt. Grass and sod may also be used as an interim control']]
        s5 = pd.DataFrame(shold, columns=tab6.columns)
        tab6 = pd.concat([tab6, s5], axis=0)

    detrmlis = []  # blank list to hold deteriorated lbp locations
    detcomlis = []  # blank list to hold deteriorated lbp components
    if xrf_pos2.shape[1] != 0:
        for index, row in xrf_pos2.iterrows():
            detrmlis.append(xrf_pos2.at[index, 'Room'])
            detcomlis.append(xrf_pos2.at[index, 'Component²'])
        shold = [['Deteriorated Lead Based Paint',
                 str(unq_lis(detrmlis)),
                 str(unq_lis(detcomlis)),
                 'Abatement, Enclosure, Encapsulation or Paint Film Stabilization']]
        s6 = pd.DataFrame(shold, columns=tab6.columns)
        tab6 = pd.concat([tab6, s6], axis=0)

    # ------------------------------------------------------------------------------------------------------------------

    dflis = {'xrf_pos1': pd.DataFrame(xrf_pos1),  # group df's into a dictionary
             'xrf_pos2': pd.DataFrame(xrf_pos2),
             'xrf_pos3': pd.DataFrame(xrf_pos3),
             'pb_wipes3': pd.DataFrame(pb_wipes3),
             'pb_drip': pd.DataFrame(pb_drip),
             'tab6': tab6}
    for x in dflis:
        dflis[x] = dflis[x].reset_index(drop=True)  # reset indices
        if dflis[x].empty:
            xdata = [{'Room': 'None Found',
                      'Side': 'N/A',
                      'Component²': 'N/A',
                      'Substrate': 'N/A',
                      'Condition': 'N/A',
                      'Color': 'N/A'}]
            dflis[x] = pd.DataFrame(xdata)

    return dflis


# input: raw xrf data as df
# output: clean xrf data as df
def xrf_cleaner(xrf_dirty):
    # ----------------------------------------------------------------------------------------------------------------------
    # input: raw xrf data
    # output: clean xrf excel sheet
    # ----------------------------------------------------------------------------------------------------------------------

    xrf = pd.DataFrame(xrf_dirty)  # save argument as variable xrf

    for index, row in xrf.iterrows():  # remove xrf model block at top left of sheet
        if xrf.loc[index][0] == 'Company':  # remove xrf info block
            if xrf.loc[index][0] != 'Job Id':
                xrf = xrf.drop(index)

    xrf = xrf.reset_index(drop=True)  # reset index
    xrf['-->Member'] = xrf['-->Member'].replace(
        ['N/A', '---'],
        ['', ''])

    for index, row in xrf.iterrows():  # concat column 23 and 24
        if str(row[24]) != 'nan':
            xrf.iloc[index, 23] = str(row[23]) + ' ' + str(row[24])
        else:
            xrf.iloc[index, 23] = str(row[23])

    xrf = xrf.iloc[:, [1, 19, 21, 22, 23, 25, 26, 28, 5, 2, 3]]  # choose columns
    xrf.columns.values[4] = 'Component'  # rename column
    xrf.columns.values[9] = 'PbC'  # rename column
    xrf['Component'] = xrf['Component'].replace(
        ['Exterior', 'Room'],
        ['', ''])

    for index, row in xrf.iterrows():  # clear calibration cells of irrelevant data
        if xrf.loc[index][2] == 'Calibration':
            xrf.loc[index, 'Side':'Color'] = ''
        if 'Room ' in row[4]:
            xrf.loc[index, 'Component'] = str(xrf.loc[index, 'Component']).replace('Room ', '')
        if 'Exterior ' in row[4]:
            xrf.loc[index, 'Component'] = str(xrf.loc[index, 'Component']).replace('Exterior ', '')

    xrf['Job'] = xrf['Job'].apply(lambda x: x.title())  # capitalize first letter of every word in "Job" column

    return xrf


# input: schedule row relevant to app #, as array beholden
# output: clean xrf df by calling xrf-cleaner() on the raw xrf file "readings" in LBP folder
def get_xrf(beholden):
    # get path of readings file in dir
    f_name = str(beholden[0]) + ' - ' + str(beholden[5] + ' - ' + str(beholden[6]))
    file_name = os.listdir('job_Folders/' + f_name + '/' + str(beholden[2]) + '_LBP/xrf_Data_Raw')
    file_path = 'job_Folders/' + f_name + '/' + str(beholden[2]) + '_LBP/xrf_Data_Raw/' + file_name[0]
    suffix = file_name[0].split('.')[1]
    if suffix == 'xlsx':
        xhold = pd.read_excel(file_path)
    if suffix == 'csv':
        xhold = pd.read_csv(file_path, skiprows=5)  # excel sheet must be named this in the folder
    try:
        return xrf_cleaner(xhold)
    except:
        print('get xrf fail')


# input: schedule containing to-do app numbers
# output: clean schedule as df
def parse_excel(schedule):
    pd.options.mode.chained_assignment = None  # default='warn'

    # concatenate pertinent excel columns
    wkbk = pd.read_excel(schedule)  # excel sheet must be named this in the folder
    wkbk_1 = wkbk.iloc[0:44, 0:9]  # concatenate the important columns in schedule
    wkbk_2 = wkbk.iloc[:, 10:12]  # bad touch
    wkbk_3 = [wkbk_1, wkbk_2.reindex(wkbk_1.index)]  # match indexes between split dataframes
    wkbk_3 = pd.concat(wkbk_3, axis=1)

    for i, r in wkbk_3.iterrows():  # expand inspector column data to fill cells top down
        if str(r[1]) != 'nan':
            row = r
        else:
            try:
                wkbk_3.Inspector[i] = row[1]
                wkbk_3.Day[i] = row[9]
                wkbk_3.Date[i] = row[10]
            except:
                pass

    # fill in empty cells in the "Inspector" column
    # name must be located in top line of block
    for index, row in wkbk_3[::-1].iterrows():
        if str(row[1]) == 'nan':
            wkbk_3.Inspector[index] = wkbk_3.Inspector[index+1]

    wkbk_3.insert(2, 'lead_Num', '')  # add blank column
    wkbk_4 = wkbk_3.rename({'Inspection #': 'acm_Num'}, axis=1)  # change name

    for index, row in wkbk_4.iterrows():  # add lead inspector #'s to lead column
        if 'lead' in str(row[4]).lower() and 'acm' in str(row[4]).lower():
            if '/' in str(row[3]):
                num_hold = str(row[3]).split('/')
            elif ',' in str(row[3]):
                num_hold = str(row[3]).split(',')
            if len(num_hold) == 2:
                lead_hold = num_hold[0].strip()
                asb_hold = num_hold[1].strip()
                wkbk_4.lead_Num[index] = lead_hold
                wkbk_4.acm_Num[index] = asb_hold
        if 'lead' in str(row[4]).lower() and 'acm' not in str(row[4]).lower():
            wkbk_4.lead_Num[index] = str(row[3])

    for index, row in wkbk_4.iterrows():  # remove lead inspector #'s from asb column
        if 'ACM' not in str(row[4]):
            wkbk_4.acm_Num[index] = ''

    return wkbk_4  # output finished dataframe


# input: pdf_path is the relative path to the lab results
# output: two variables, dust wipes and soil; as dfs
def pdf_scrape(pdf_path):
    pb_res_df = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)  # read table_Example.pdf into dataframe
    pb_dripline_df = tabula.read_pdf(pdf_path, pages=3, area=[204.691, 34.689, 262.846, 566.503])

    return pb_res_df, pb_dripline_df


# input: xrf is the output of xrf-cleaner()
# input: beholden is an array containing the row of schedule relavant to the app #
# output: clean xrf df saved as excel file
def save_xrf_clean_xlsx(xrf, beholden):
    app_folder = r'C:\Users\Elliott\pythonplay\lead_Pit\LRA\finished_Docs'
    appf = app_folder + '\\' + str(beholden[0])
    if not os.path.exists(appf):
        os.makedirs(appf)

    filename_xrf = 'lead_Pit/LRA/finished_Docs/' + str(beholden[0]) + '/xrf_clean.xlsx'
    writer = pd.ExcelWriter(filename_xrf, engine='xlsxwriter')
    workbook_setter = writer.book  # add formatting function
    header_format = workbook_setter.add_format({'bold': True,
                                                'text_wrap': False,
                                                'fg_color': '348feb',
                                                'border': 1,
                                                'align': 'center'})  # header format
    body_format = workbook_setter.add_format({'text_wrap': False,
                                              'border': 1,
                                              'align': 'center'})  # body format
    positive_format = workbook_setter.add_format({'bg_color': 'ff1a1a',
                                                  'text_wrap': False,
                                                  'border': 1,
                                                  'align': 'center'})
    sht_nm = 'xrf_data'
    xrf.to_excel(writer, sheet_name=sht_nm, startrow=0, index=False)
    sheet = writer.sheets[sht_nm]

    for x in range(len(xrf.columns)):  # create blank sheet with labeled columns
        loc_str = string.ascii_uppercase[x] + str(1)
        sheet.write(loc_str, xrf.columns[x], header_format)

    for i in range(2, xrf.shape[0]+2):  # row
        for j in range(xrf.shape[1]):  # column
            loc_str = string.ascii_uppercase[j] + str(i)
            if xrf.iat[i-2, 8] == 'Positive' and xrf.iat[i-2, 2] != 'Calibration':
                sheet.write(loc_str, xrf.iat[i-2, j], positive_format)
            else:
                try:
                    sheet.write(loc_str, xrf.iat[i-2, j], body_format)
                except:
                    sheet.write(loc_str, '---', body_format)
    writer.save()


# input: excel file xrf_clean
# output: clean xrf data saved as xrf_clean.pdf
def xrf_clean_excel2pdf(xrf, beholden):

    app_folder = r'C:\Users\Elliott\pythonplay\lead_Pit\LRA\finished_Docs'
    appf = app_folder + '\\' + str(beholden[0])
    appff = appf + '\\' + 'xrf_clean.pdf'

    colors = []
    deff = '#FFFFFF'
    poss = '#ff1a1a'
    coll = '#3380FF'
    deffrow = []
    possrow = []
    cols = []
    for x in range(xrf.shape[1]):
        deffrow.append(deff)
        possrow.append(poss)
        cols.append(coll)

    for index, row in xrf.iterrows():
        if row[8] == 'Positive' and row[2] != 'Calibration':
            colors.append(possrow)
        else:
            colors.append(deffrow)
    fig, ax = plt.subplots(figsize=(12, 4))
    ax.axis('tight')
    ax.axis('off')
    the_table = ax.table(cellText=xrf.values,
                         colLabels=xrf.columns,
                         loc='center',
                         cellLoc='center',
                         colColours=cols,
                         cellColours=colors)
    the_table.auto_set_column_width(col=list(range(len(xrf.columns))))
    pp = PdfPages(str(appff))
    pp.savefig(fig, bbox_inches='tight')
    pp.close()


# input: xrf is the output of xrf-cleaner()
# input: beholden is an array containing the row of schedule relavant to the app #
# output: clean xrf df saved as excel file containing tables 1-5 as separate sheets named 'table1', 'table2', etc.
def save_xrf_pos_xlsx(dflis, beholden):
    table_names = ['Table 1: Lead-Based Paint¹',
                   'Table 2: Deteriorated Lead-Based Paint¹',
                   'Table 3: Lead Containing Materials²',
                   'Table 4: Dust Wipe Sample Analysis',
                   'Table 5: Soil Sample Analysis',
                   'Table 6: Lead Hazard Control Options¹']
    # create excel file to hold the different tables as separate worksheets
    filename_xrf = 'lead_Pit/LRA/finished_Docs/' + str(beholden[0]) + '/xrf_pos.xlsx'
    writer = pd.ExcelWriter(filename_xrf, engine='xlsxwriter')

    workbook_setter = writer.book  # add formatting function
    header_format = workbook_setter.add_format({'bold': True,
                                                'text_wrap': False,
                                                'fg_color': '348feb',
                                                'border': 1,
                                                'align': 'center'})  # header format
    body_format = workbook_setter.add_format({'text_wrap': False,
                                              'border': 1,
                                              'align': 'center'})  # body format

    for x, df in enumerate(dflis):  # create sheets
        sht_nm = 'table' + str(x+1)
        dflis[df].to_excel(writer, sheet_name=sht_nm, startrow=0, index=False)
        sheet = writer.sheets[sht_nm]
        if x < 3:
            sheet.merge_range('A2:F2', table_names[x], header_format)
        elif x == 5:
            sheet.merge_range('A2:D2', table_names[x], header_format)
        else:
            sheet.merge_range('A2:E2', table_names[x], header_format)
        for y in range(len(dflis[df].columns)):
            loc_str = string.ascii_uppercase[y] + str(3)
            sheet.write(loc_str, dflis[df].columns[y], header_format)

        # write to sheets and apply body format to body
        for i in range(4, dflis[df].shape[0]+4):  # number value
            for j in range(dflis[df].shape[1]):  # letter value
                loc_str = string.ascii_uppercase[j] + str(i)
                try:
                    sheet.write(loc_str, dflis[df].iat[i-4, j], body_format)
                except:
                    sheet.write(loc_str, '---', body_format)

        sheet.set_row(0, None, None, {'hidden': True})  # hide original header row
    writer.save()


def create_lra(dfliss, beholden, insp_num, proj_no):
    # make list of lengths of tables stored in dflis for doc formatting
    dflis = list(dfliss.values())
    dflis_lenst = []
    for x in range(len(dflis)):
        dflis_lenst.append(dflis[x].shape[0])

    # set up blank word document named "doc"
    doc = docx.Document()  # create instance of a word document
    sections = doc.sections  # change the page margins
    for section in sections:  # set margins equal to 0 on all sides of doc container
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

    # ------------------------------------------------------------------------------------------------------------------
    # PAGE 1
    # ------------------------------------------------------------------------------------------------------------------

    doc.add_paragraph('')
    doc.add_picture('lead_Pit/LRA/LRA_Header.jpg', width=Inches(8.5))  # add image with defined size

    para = doc.add_paragraph(datetime.now().strftime('%m/%d/%y'))  # add today's date

    para.add_run('\n\nDewberry c/o NCORR\n'
                 '1545 Peachtree Street NE, Suite 250\n'
                 'Atlanta, Georgia 30309\n\n'
                 'Re:      Lead Risk Assessment\n'
                 '            ' + str(beholden[5]) + ', ' + str(beholden[6]) + ', NC ' + str(beholden[7]) + '\n'
                 '            EI Project No: IHMO' + proj_no + '\n\n')

    para.add_run('Project Site Address: ').bold = True
    para.add_run(str(beholden[5]) + ', ' + str(beholden[6]) + ', NC ' + str(beholden[7]))

    para.add_run('\n\nNCORR APP ID: ').bold = True
    para.add_run(beholden[0] + ', ' + str(beholden[2]))

    para.add_run('\n\nInspection Date: ').bold = True
    para.add_run(beholden[11].strftime('%m/%d/%y'))

    para.add_run('\n\nScope of Work: ').bold = True
    para.add_run('Lead Risk Assessment')

    para.add_run('\n\nLead-Based Paint Inspection: ').bold = True  # touch
    if dflis[0].shape[0] > 0:
        para.add_run('Lead-Based Paint Found')
    else:
        para.add_run('Lead-Based Paint Not Found')

    para.add_run('\n\nDeteriorated Lead-Based Paint: ').bold = True
    if dflis[1].shape[0] > 0:
        para.add_run('Yes')
    else:
        para.add_run('No')

    para.add_run('\n\nLead Containing Materials: ').bold = True
    if dflis[2].shape[0] > 0:
        para.add_run('Yes')
    else:
        para.add_run('No')

    para.add_run('\n\nLead Dust Hazards: ').bold = True
    if 'Yes' in dflis[3].iloc[:, 4].values:
        para.add_run('Yes')
    else:
        para.add_run('None Found')

    para.add_run('\n\nLead Soil Hazards: ').bold = True
    if 'Yes' in dflis[4].iloc[:, 4].values:
        para.add_run('Yes')
    else:
        para.add_run('None Found')

    para.add_run('\n\nRecommendations: ').bold = True
    para.add_run('Recommendations for lead-based paint hazards: see Table 6')

    para.add_run('\n\nInspector: ').bold = True
    para.add_run(search_arr(insp_num, beholden[1])[0] + ', North Carolina Risk Assessor #' + search_arr(insp_num,
                                                                                                        beholden[1])[1])

    # add section for 2 inspector names and signatures at the bottom of the first page
    emp_name = beholden[1]  # assign current employee name to variable emp_name
    emp_lis = os.listdir('lead_Pit/reporting/sig_Block')  # create list of employees from sig_Block folder

    # create new doc element 'new_para' to hold signature block
    for x in range(len(emp_lis)):
        if emp_name.lower() in emp_lis[x]:
            sig_path = emp_lis[x]  # assign file path of current employee sig block to variable sig_path
            sig_path = 'lead_Pit/reporting/sig_Block/' + sig_path
            doc.add_picture(sig_path)
            new_para = doc.paragraphs[-1]
            new_para.paragraph_format.left_indent = Inches(0.85)

    doc.add_picture('lead_Pit/LRA/LRA_Footer.jpg', width=Inches(8.5))  # add footer

    # ------------------------------------------------------------------------------------------------------------------
    # PAGE 2
    # ------------------------------------------------------------------------------------------------------------------

    notes_style = doc.styles
    notes_charstyle = notes_style.add_style('NotesStyle', WD_STYLE_TYPE.CHARACTER)
    notes_font = notes_charstyle.font
    notes_font.size = Pt(9)
    notes_font.name = 'Arial'

    para1 = doc.add_paragraph('\n')
    para1.add_run('\n\n1.  Findings: \n').bold = True  # header of table 1

    t1 = doc.add_table(dflis[0].shape[0] + 2, dflis[0].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t1, 0)  # populate table 1

    para2 = doc.add_paragraph('\n')
    run1 = para2.add_run('Note(s):\n'
                         '     1.  Positive results indicate lead in quantities equal to or greater than 1.0 mg/cm² and are considered lead-based\n          paint.\n'
                         '     2.  Samples are taken to represent component types; therefore, it should be assumed that similar component\n          types in the rest of that room of room equivalent also contain lead-based paint.\n',
                         style='NotesStyle')

    runhold = '0'
    if dflis_lenst[0] > 20:
        run1.add_break(WD_BREAK.PAGE)
        para2.add_run('\n\n')
        runhold = '1'

    t2 = doc.add_table(dflis[1].shape[0] + 2, dflis[1].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t2, 1)  # populate table 2

    para3 = doc.add_paragraph('\n')
    run2 = para3.add_run('Note(s):\n' + '     1.  Surfaces in deteriorated condition are considered to be lead-based paint hazards as defined by Title X and\n          should be addressed through abatement or interim controls which are described in Table 6.\n',
                         style='NotesStyle')

    if runhold == '0' and (dflis_lenst[0] + dflis_lenst[1]) > 15:
        run2.add_break(WD_BREAK.PAGE)
        runhold = '2'

    t3 = doc.add_table(dflis[2].shape[0] + 2, dflis[2].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t3, 2)  # populate table 3

    para4 = doc.add_paragraph('\n')
    run3 = para4.add_run('Note(s):\n' + '     2.  Although not considered to be lead-based paint, these materials when disturbed through destructive measures\n          such as sanding, chipping, grinding, and other sourceds of friction, can create dust hazards and should be\n          treated through control described in Table 6.\n',
                         style='NotesStyle')

    if runhold == '0':
        run3.add_break(WD_BREAK.PAGE)
        para4.add_run('\n\n')

    # ------------------------------------------------------------------------------------------------------------------
    # PAGE 3
    # ------------------------------------------------------------------------------------------------------------------

    t4 = doc.add_table(dflis[3].shape[0] + 2, dflis[3].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t4, 3)  # populate table 4

    para5 = doc.add_paragraph('\n')
    run5 = para5.add_run('Note(s):\n'
                         '     1.  EPA Lead Dust Hazard for Floors: 10 μg/ft²\n',
                         style='NotesStyle')

    if runhold == '1':
        run5.add_break(WD_BREAK.PAGE)
        para5.add_run('\n\n')

    t5 = doc.add_table(dflis[4].shape[0] + 2, dflis[4].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t5, 4)  # populate table 5

    para6 = doc.add_paragraph('\n')
    run6 = para6.add_run('Note(s):\n'
                  '     1.  EPA Lead in Soil Hazard for children\'s play areas with bare residential soil: 400 mg/Kg; bare soil for the\n          remainder of the yard: 1,200 mg/Kg\n'
                  '_____________________________________________________________________________________________',
                  style='NotesStyle')

    if runhold == '0':
        run6.add_break(WD_BREAK.PAGE)
        para6.add_run('\n')

    # ------------------------------------------------------------------------------------------------------------------
    # PAGE 4
    # ------------------------------------------------------------------------------------------------------------------

    para7 = doc.add_paragraph()
    para7.add_run('\n')

    t6 = doc.add_table(dflis[5].shape[0]+2, dflis[5].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t6, 5)  # populate table 6

    para8 = doc.add_paragraph('\n')
    run8 = para8.add_run('Note(s):\n'
                  '     1.  Lead hazard control options include abatement and interim controls.\n'
                  '     2.  Paint film stabilization: Wet scrape and prime building comopnents where chipping or peeling is present\n          following acceptable methods.\n'
                  '     3.  Replace: Remove and dispose of components in accordance with applicable federal, state and local\n          regulations. Prime coat any new unpainted wood components.\n'
                  '     4.  Enclosure: Enclose lead-based paint coated building components with a material that is structurally affixed and\n          deemed to last 20 years.\n'
                  '     5.  General Cleaning-Clean using HEPA filtered vacuum and wet wipe impacted surfaces to remove paint chips\n          and lead-dust hazards.\n'
                  '    _____________________________________________________________________________________________',
                  style='NotesStyle')

    if runhold == '1':
        run8.add_break(WD_BREAK.PAGE)
        para8.add_run('\n')

    para8.add_run('\n\n\n2.  Limitations: \n\n').bold = True  # header of table 1
    para8.add_run('    •   No limitations were encountered during the course of this survey\n'
                  '    •   Exterior windows were inaccessible due to storm window coverings\n'
                  '    •   No soil was observed along the dripline, therefore no soil sample was collected\n')

    para8.add_run('\n\n3.  Lead Hazard Control Activities:\n').bold = True  # header of table 1

    par1 = doc.add_paragraph()
    par1.add_run('All lead abatement activities must be performed in strict compliance with the Department of Housing and Urban Development (HUD) 24 CFR Part 35, and the Environmental Protection Agency')

    para9 = doc.add_paragraph('(EPA) 40 CFR Part 745 Subpart L.\n')
    para9.add_run()

    par2 = doc.add_paragraph()
    par2.add_run('All contractor’s personnel who will disturb lead-based paint during the course of their work on this residence should be informed of the potential danger posed by lead-based paint and should be')

    para10 = doc.add_paragraph()
    para10.add_run('directed to comply with all applicable federal, state, and local lead abatement regulations.\n')

    par3 = doc.add_paragraph()
    par3.add_run('Table 6 lists each lead hazard identified, along with control options. Highest priority should be given to correcting lead hazards with greater probability of being contacted by children six years of age and under, women who are or may become pregnant, and residents of the home. These include, but are not limited to, deteriorated lead-based paint inside the residence on friction and impact surfaces (windows and doors), other surfaces (i.e. walls or trims) at a height of six feet and below, lead dust hazards, deteriorated lead-based paint on exterior friction and impact surfaces (windows')

    para11 = doc.add_paragraph()
    run11 = para11.add_run('and doors), and lead soil hazards in children’s play areas.\n')

    if runhold == '0':
        run11.add_break(WD_BREAK.PAGE)
        para11.add_run('\n\n')

    # ------------------------------------------------------------------------------------------------------------------
    # PAGE 5
    # ------------------------------------------------------------------------------------------------------------------

    par4 = doc.add_paragraph()
    par4.add_run('If paint condition is intact, no treatment is required at this time. However, ongoing monitoring and maintenance of painted surfaces containing lead-based paint must be performed on a routine basis as paint conditions may deteriorate potentially creating a lead dust hazard. Painted surfaces should be inspected annually and repainted as needed before deterioration occurs. Prior to any scraping or sanding, appropriate measures should be taken to prevent the generation or spreading of paint')

    para12 = doc.add_paragraph()
    para12.add_run('chips or dust.\n\n')

    para13 = doc.add_paragraph('')
    para13.add_run('4.  HUD Notification: \n\n').bold = True  # header of table 1
    run13 = para13.add_run('A copy of this summary must be provided to new lessees (tenants) and purchasers of this property under Federal Law (24 CFR part 35 and 40 CFR part 745) before they become obligated under a lease or sales contract. The complete report must also be provided to new purchasers and be made available to new tenants. Landlords (lessors) and sellers are also required to distribute an educational pamphlet and include standard warning language in their leases or sales contracts, to ensure that parents have the information necessary to protect their children from lead-based paint hazards.')

    if runhold == '1':
        run13.add_break(WD_BREAK.PAGE)
        para13.add_run('\n\n')

    para13.add_run('\n\n\n3.  Lead Hazard Control Activities:\n\n').bold = True  # header of table 1
    para13.add_run('    •    Floor Plan/Diagram\n'
                   '    •    Risk Assessment Forms\n'
                   '    •    XRF Data Sheets/Photo Log\n'
                   '    •    Lab Results/Chain of Custody\n'
                   '    •    Methodology\n'
                   '    •    Lead Hazard Control Options\n'
                   '    •    Definitions\n'
                   '    •    Lead Based Paint Activity Summary (LBPAS)\n'
                   '    •    XRF Analyzer Performance Characteristics Sheet\n'
                   '    •    Certifications and Licensure').bold = True

    # ------------------------------------------------------------------------------------------------------------------
    # FORMATTING

    para.paragraph_format.left_indent = Inches(0.85)  # indent paragraphs above
    para.paragraph_format.line_spacing = 1.03  # line height
    for x in range(1, 14):  # left pad all note paragraphs 0.9 inches
        try:
            pholdd = 'para' + str(x)
            eval(pholdd).paragraph_format.left_indent = Inches(0.9)
            eval(pholdd).paragraph_format.right_indent = Inches(1.1)
            eval(pholdd).paragraph_format.space_before = Inches(0)
            eval(pholdd).paragraph_format.space_after = Inches(0)
        except:
            pass
    for x in range(1, 10):
        try:
            pholdd1 = 'par' + str(x)
            eval(pholdd1).paragraph_format.left_indent = Inches(0.9)
            eval(pholdd1).paragraph_format.right_indent = Inches(1.1)
            eval(pholdd1).paragraph_format.space_before = Inches(0)
            eval(pholdd1).paragraph_format.space_after = Inches(0)
            eval(pholdd1).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except:
            pass

    # ------------------------------------------------------------------------------------------------------------------
    # SAVE DOCUMENT AS

    # create new app folder containing LRA word doc
    app_folder = r'C:\Users\Elliott\pythonplay\lead_Pit\LRA\finished_Docs'
    appf = app_folder + '\\' + str(beholden[0])
    appff = appf + '\\' + str(beholden[0]) + '_LRA.docx'
    if not os.path.exists(appf):
        os.makedirs(appf)
    doc.save(str(appff))  # save doc


def create_lbpas(dfliss, beholden, insp_num, sig):
    dflis = list(dfliss.values())
    dflis_lenst = []
    for x in range(len(dflis)):
        dflis_lenst.append(dflis[x].shape[0])

    insp_names = insp_num.keys()
    name_hold = ''
    for name in insp_names:
        if beholden[1] in name:
            name_hold = name

    sig_pat = 'lead_Pit/reporting/Signatures/' + sig[name_hold] + '.png'
    doc = docx.Document()  # create instance of a word document
    sec_pr = doc.sections[0]._sectPr  # get the section properties el
    pg_borders = OxmlElement('w:pgBorders')  # create new borders el
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')  # a single line
        border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el)  # register single border to border el
    sec_pr.append(pg_borders)  # apply border changes to section

    sections = doc.sections  # change the page margins
    for section in sections:  # set margins equal to 0 on all sides of doc container
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.5)

    # top left text
    tl_text = doc.styles
    tl_charstyle = tl_text.add_style('TLStyle', WD_STYLE_TYPE.CHARACTER)
    tl_font = tl_charstyle.font
    tl_font.size = Pt(6)
    tl_font.name = 'Arial'

    # tiny text
    tiny_text = doc.styles
    notes_charstyle = tiny_text.add_style('TinyStyle', WD_STYLE_TYPE.CHARACTER)
    tiny_font = notes_charstyle.font
    tiny_font.size = Pt(9)
    tiny_font.name = 'Arial'
    tiny_font.bold = True

    # header
    head_text = doc.styles
    head_charstyle = head_text.add_style('HeadStyle', WD_STYLE_TYPE.CHARACTER)
    head_font = head_charstyle.font
    head_font.size = Pt(13)
    head_font.name = 'Arial'
    head_font.bold = True
    head_font.underline = True

    # subheader
    subhead_text = doc.styles
    subhead_charstyle = subhead_text.add_style('SubHeadStyle', WD_STYLE_TYPE.CHARACTER)
    subhead_font = subhead_charstyle.font
    subhead_font.size = Pt(8)
    subhead_font.name = 'Arial'

    # roman numerals
    roman_text = doc.styles
    roman_charstyle = roman_text.add_style('RomanStyle', WD_STYLE_TYPE.CHARACTER)
    roman_font = roman_charstyle.font
    roman_font.size = Pt(12)
    roman_font.name = 'Arial'
    roman_font.bold = True

    # regular forms
    reg_form_text = doc.styles
    reg_form_charstyle = reg_form_text.add_style('RegFormStyle', WD_STYLE_TYPE.CHARACTER)
    reg_form_font = reg_form_charstyle.font
    reg_form_font.size = Pt(9)
    reg_form_font.name = 'Arial'
    reg_form_font.bold = True

    # underlined forms
    form_text = doc.styles
    form_charstyle = form_text.add_style('FormStyle', WD_STYLE_TYPE.CHARACTER)
    form_font = form_charstyle.font
    form_font.size = Pt(9)
    form_font.name = 'Arial'
    form_font.bold = True
    form_font.underline = True

    # notes style
    notes_style = doc.styles
    notes_charstyle = notes_style.add_style('NotesStyle', WD_STYLE_TYPE.CHARACTER)
    notes_font = notes_charstyle.font
    notes_font.size = Pt(9)
    notes_font.name = 'Arial'

    para = doc.add_paragraph('')
    para.add_run('NC DEPARTMENT OF HEALTH AND HUMAN SERVICES\nDIVISION OF PUBLIC HEALTH\nHEALTH HAZARDS CONTROL UNIT\n',
                 style='TLStyle')

    para1 = doc.add_paragraph()
    para1.add_run('LEAD-BASED PAINT ACTIVITY SUMMARY\n',
                 style='HeadStyle')
    para1.add_run('**Please type or print in ink.**',
                  style='SubHeadStyle')

    para3 = doc.add_paragraph()

    para3.add_run('I.   ',
                  style='RomanStyle')
    para3.add_run('TYPE OF ACTIVITY:',
                  style='RegFormStyle')
    para3.add_run('\n\n         _____ Inspection          __',
                  style='TinyStyle')
    para3.add_run('x',
                  style='FormStyle').bold = False

    para3.add_run('__ Risk Assessment          _____ Lead Hazard Screen',
                  style='TinyStyle')

    para3.add_run('\n\nII.   ',
                  style='RomanStyle')
    para3.add_run('DATE ACTIVITY COMPLETED: ' + beholden[11].strftime('%B %d, %y'),
                  style='TinyStyle')

    para3.add_run('\n\nIII.   ',
                  style='RomanStyle')
    para3.add_run('ACTIVITY LOCATION: ' + beholden[5],
                  style='TinyStyle')

    para3.add_run('\n\n         Address:  ',
                  style='TinyStyle')
    para3.add_run('  ' + beholden[5] + '  ',
                  style='FormStyle')

    para3.add_run('\n\n         City:  ',
                  style='TinyStyle')
    para3.add_run('  ' + beholden[6] + '  ',
                  style='FormStyle')
    para3.add_run('   State:  ',
                  style='TinyStyle')
    para3.add_run('  NC  ',
                  style='FormStyle')
    para3.add_run('   Zip Code:  ',
                  style='TinyStyle')
    para3.add_run('  ' + str(beholden[7]) + '  ',
                  style='FormStyle')
    para3.add_run('   County:  ',
                  style='TinyStyle')
    para3.add_run('  ' + str(beholden[8]) + '  ',
                  style='FormStyle')

    para3.add_run('\n\n         Contact Person:  ',
                  style='TinyStyle')
    para3.add_run('   Charles Aly   ',
                  style='FormStyle')
    para3.add_run('   Contact Phone:  ',
                  style='TinyStyle')
    para3.add_run('  678-205-6903  ',
                  style='FormStyle')

    para3.add_run('\n\nIV.   ',
                  style='RomanStyle')
    para3.add_run('ACTIVITY SUMMARY (attach additional pages as needed):',
                  style='TinyStyle')

    t1 = doc.add_table(dflis[0].shape[0] + 2, dflis[0].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t1, 0)  # populate table 1

    par0 = doc.add_paragraph('\n')
    run0 = par0.add_run('Note(s):\n'
                         '     1.  Positive results indicate lead in quantities equal to or greater than 1.0 mg/cm² and are considered lead-based\n          paint.\n'
                         '     2.  Samples are taken to represent component types; therefore, it should be assumed that similar component\n          types in the rest of that room of room equivalent also contain lead-based paint.\n',
                         style='NotesStyle')

    runhold = 0
    if dflis_lenst[0] > 6:
        run0.add_break(WD_BREAK.PAGE)
        runhold = 1

    t2 = doc.add_table(dflis[1].shape[0] + 2, dflis[1].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t2, 1)  # populate table 2

    par1 = doc.add_paragraph('\n')
    run1 = par1.add_run('Note(s):\n' + '     1.  Surfaces in deteriorated condition are considered to be lead-based paint hazards as defined by Title X and\n          should be addressed through abatement or interim controls which are described in Table 6.\n',
                         style='NotesStyle')

    if runhold == 0:
        run1.add_break(WD_BREAK.PAGE)
        runhold = 2

    t3 = doc.add_table(dflis[2].shape[0] + 2, dflis[2].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t3, 2)  # populate table 3

    par2 = doc.add_paragraph('\n')
    run2 = par2.add_run('Note(s):\n' + '     2.  Although not considered to be lead-based paint, these materials when disturbed through destructive measures\n          such as sanding, chipping, grinding, and other sourceds of friction, can create dust hazards and should be\n          treated through control described in Table 6.\n',
                         style='NotesStyle')

    set1 = int(dflis_lenst[1] + dflis_lenst[2])
    if runhold == 1 and set1 > 8:
        run2.add_break(WD_BREAK.PAGE)
        runhold = 2

    if runhold == 2 and dflis_lenst[2] > 8:
        run2.add_break(WD_BREAK.PAGE)
        runhold = 2

    t4 = doc.add_table(dflis[3].shape[0] + 2, dflis[3].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t4, 3)  # populate table 4

    par3 = doc.add_paragraph('\n')
    run3 = par3.add_run('Note(s):\n'
                         '     1.  EPA Lead Dust Hazard for Floors: 10 μg/ft²\n',
                         style='NotesStyle')

    if runhold == 2:
        run3.add_break(WD_BREAK.PAGE)
        runhold = 3

    t5 = doc.add_table(dflis[4].shape[0] + 2, dflis[4].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t5, 4)  # populate table 5

    par4 = doc.add_paragraph('\n')
    run4 = par4.add_run('Note(s):\n'
                  '     1.  EPA Lead in Soil Hazard for children\'s play areas with bare residential soil: 400 mg/Kg; bare soil for the\n          remainder of the yard: 1,200 mg/Kg\n'
                  '_____________________________________________________________________________________________\n',
                  style='NotesStyle')

    if runhold == 1 or runhold == 2:
        run4.add_break(WD_BREAK.PAGE)
        runhold == 4

    t6 = doc.add_table(dflis[5].shape[0]+2, dflis[5].shape[1])  # add blank table to doc using shape of dflis
    pop_table(dflis, t6, 5)  # populate table 6

    par5 = doc.add_paragraph('\n')
    run5 = par5.add_run('Note(s):\n'
                  '     1.  Lead hazard control options include abatement and interim controls.\n'
                  '     2.  Paint film stabilization: Wet scrape and prime building comopnents where chipping or peeling is present\n          following acceptable methods.\n'
                  '     3.  Replace: Remove and dispose of components in accordance with applicable federal, state and local\n          regulations. Prime coat any new unpainted wood components.\n'
                  '     4.  Enclosure: Enclose lead-based paint coated building components with a material that is structurally affixed and\n          deemed to last 20 years.\n'
                  '     5.  General Cleaning-Clean using HEPA filtered vacuum and wet wipe impacted surfaces to remove paint chips\n          and lead-dust hazards.\n'
                  '    __________________________________________________________________________________________',
                  style='NotesStyle')

    par6 = doc.add_paragraph()
    run9 = par6.add_run()

    if runhold == 3:
        run9.add_break(WD_BREAK.PAGE)
        runhold = 5

    para6 = doc.add_paragraph()
    para6.add_run('\n\nV.   ',
                  style='RomanStyle')
    para6.add_run('CERTIFIED INSPECTOR OR RISK ASSESSOR',
                  style='TinyStyle')

    para6.add_run('\n\n         Name:  ',
                  style='TinyStyle')
    para6.add_run('      ' + name_hold + '                           ',
                  style='FormStyle')
    para6.add_run('   NC Lead Cert No.:  ',
                  style='TinyStyle')
    para6.add_run('  ' + insp_num[name_hold] + '                                       _',
                  style='FormStyle')

    para6.add_run('\n\n         Title:  ',
                  style='TinyStyle')
    para6.add_run('      Industrial Hygienist                                                       _',
                  style='FormStyle')

    para6.add_run('\n\n         Certified Firm:  ',
                  style='TinyStyle')
    para6.add_run('      The EI Group, Inc                        ',
                  style='FormStyle')
    para6.add_run('   NC Cert. No:  ',
                  style='TinyStyle')
    para6.add_run('  FPB-OO18                            _',
                  style='FormStyle')

    para6.add_run('\n\n         Address:  ',
                  style='TinyStyle')
    para6.add_run('      2101 Gateway Centre Blcd. Suite 200           ',
                  style='FormStyle')
    para6.add_run('   State:  ',
                  style='TinyStyle')
    para6.add_run('  NC     ',
                  style='FormStyle')
    para6.add_run('   Zip:  ',
                  style='TinyStyle')
    para6.add_run('  27560                _',
                  style='FormStyle')

    para6.add_run('\n\n         Telephone:  ',
                  style='TinyStyle')
    para6.add_run('      919-657-7500                                                                                      _',
                  style='FormStyle')

    run6 = para6.add_run('\n\n         Signature:  ',
                  style='TinyStyle')
    run7 = para6.add_run('  ',
                         style='FormStyle')
    inline_pic0 = run7.add_picture(sig_pat, width=Inches(2))
    run7 = para6.add_run('                                      ',
                         style='FormStyle')
    para6.add_run('Date:  ',
                  style='TinyStyle')
    para6.add_run('  ' + beholden[11].strftime('%B %d, %y') + '         _',
                  style='FormStyle')
    para7 = doc.add_paragraph()
    para7.add_run('\n\n')
    para7.add_run('SUBMIT TO:      NC DHHS - HEALTH HAZARDS CONTROL UNIT\n'
                  '                      1912 MAIL SERVICE CENTER\n'
                  '                      RALEIGH, NC 27699-1912',
                  style='RomanStyle')

    para7.add_run('\n\n')
    para7.add_run('Lead-Based Paint Activity Summary(8/05; 7/07)\n'
                  'Health Hazards Control Unit',
                  style='TLStyle').bold = False

    para8 = doc.add_paragraph()
    run8 = para8.add_run()
    run8.add_break(WD_BREAK.PAGE)

    # last page bold
    last_bold_text = doc.styles
    last_bold_charstyle = last_bold_text.add_style('LastBoldStyle', WD_STYLE_TYPE.CHARACTER)
    last_bold_font = last_bold_charstyle.font
    last_bold_font.size = Pt(11)
    last_bold_font.name = 'Arial'
    last_bold_font.bold = True

    # last page not bold
    last_reg_text = doc.styles
    last_reg_charstyle = last_reg_text.add_style('LastRegStyle', WD_STYLE_TYPE.CHARACTER)
    last_bold_font = last_reg_charstyle.font
    last_bold_font.size = Pt(11)
    last_bold_font.name = 'Arial'

    para8.add_run('NC DEPARTMENT OF HEALTH AND HUMAN SERVICES\n'
                  'DIVISION OF PUBLIC HEALTH\n'
                  'HEALTH HAZARDS CONTROL UNIT',
                  style='TLStyle')

    para9 = doc.add_paragraph()
    para9.add_run('\n\nINSTRUCTIONS',
                  style='LastBoldStyle')
    para9.add_run('\nFOR COMPLETION OF LEAD-BASED PAINT ACTIVITY SUMMARY',
                  style='LastRegStyle')

    para10 = doc.add_paragraph()
    para10.add_run('\nPURPOSE',
                  style='LastBoldStyle')
    para10.add_run('\nA Lead-Based Paint Activity Summary shall be submitted to the North Carolina Lead-Based Paint Hazard Management Program (LHMP) by the certified inspector or risk assessor for each inspection, risk assessment, or lead hazard screen conducted within 45 days of the activity on a form provided or approved by the Program per LHMP Rule 10A NCAC 41C .0807(b).',
                  style='LastRegStyle')
    para10.add_run('\n\nPREPARATION',
                  style='LastBoldStyle')
    para10.add_run('\nAll information is to be filled out completely, typed or printed in ink.  Pencil is not acceptable.  Attachments are also to be typed or printed in ink. ',
                  style='LastRegStyle')

    para10.add_run('\n\nINSTRUCTIONS',
                  style='LastBoldStyle')
    para10.add_run('\n\nI.	Indicate the type of activity that was conducted.',
                  style='LastRegStyle')
    para10.add_run('\n\nII.	Enter the date the activity was completed.',
                  style='LastRegStyle')
    para10.add_run('\n\nIII.	Enter complete information about the facility where the activity occurred, including facility name,\n           address, city, state, zip code, county, the name of the facility contact, and the contact’s telephone\n           number, including area code.',
                  style='LastRegStyle')
    para10.add_run('\n\nIV.	Summarize the activities that were conducted at the site, including the results of the inspection, risk  	assessment, or lead hazard screen, and any recommendations resulting from the activity.',
                  style='LastRegStyle')
    para10.add_run('\n\n')
    para10.add_run('V.	Enter the name, NC Lead Certification Number, and title of the individual conducting the activity.',
                  style='LastRegStyle')
    para10.add_run('\n\n')
    para10.add_run('           Enter the name of the NC Certified Firm, the NC Firm Certification Number, the firm’s address, state,           zip code, and telephone number, including area code.',
                  style='LastRegStyle')
    para10.add_run('\n\n')
    para10.add_run('           Enter the original signature of the inspector or risk assessor who conducted the activity and the date           the Lead-Based Paint Activity Summary was signed.',
                  style='LastRegStyle')

    para10.add_run('\n\nCompleted Activity Summary with any attachments should be mailed to:\n\n'
                  '    NC Department of Health and Human Services\n'
                  '    Health Hazards Control Unit\n'
                  '    1912 Mail Service Center\n'
                  '    (919)707-5950\n\n'
                  'For Overnight/Express Mail:\n\n'
                  '    NC Department of Health and Human Services\n'
                  '    Health Hazards Control Unit\n'
                  '    5505 Six Forks Rd, 2nd Floor, Room D-1'
                  '    Raleigh, NC 27609',
                  style='TinyStyle')

    para10.add_run('\n\nLead-Based Paint Activity Summary(8/05; 7/07)\n'
                  'Health Hazards Control Unit',
                  style='TLStyle')

    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para7.paragraph_format.space_after = Inches(0)
    para9.paragraph_format.space_after = Inches(0)
    para8.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para9.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for j in range(6):
        try:
            parho = 'par' + str(j)
            eval(parho).paragraph_format.left_indent = Inches(0.5)
            eval(parho).paragraph_format.right_indent = Inches(0.1)
            eval(parho).paragraph_format.space_before = Inches(0)
            eval(parho).paragraph_format.space_after = Inches(0)
        except:
            pass

    # ------------------------------------------------------------------------------------------------------------------
    # SAVE DOCUMENT AS

    # create new app folder containing LRA word doc
    app_folder = r'C:/Users/Elliott/pythonplay/lead_Pit/LRA/finished_Docs/'
    appf = str(app_folder) + str(beholden[0]) + '/'
    appff = appf + str(beholden[0]) + '_LBPAS.docx'
    if not os.path.exists(appf):
        os.makedirs(appf)
    doc.save(str(appff))  # save doc

